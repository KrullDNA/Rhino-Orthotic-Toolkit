# -*- coding: utf-8 -*-
"""Convert the BEGINNERS_GUIDE markdown files into .docx using python-docx.

Run from repo root:
    python3 Grasshopper/_build_docs.py
"""
import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = os.path.dirname(os.path.abspath(__file__))


def add_table_from_md(doc, header_cells, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        hdr[i].text = h.strip()
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows, start=1):
        rc = table.rows[ri].cells
        for i, cell in enumerate(row):
            if i < len(rc):
                rc[i].text = cell.strip()


def add_inline(paragraph, text):
    """Render bold (**x**) and code (`x`) inside a paragraph."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def md_to_docx(md_path, docx_path, title, font_name="Calibri"):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)

    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.name = font_name

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.name = font_name
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.name = font_name
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.name = font_name
        elif line.startswith("---"):
            doc.add_paragraph().add_run("─" * 40)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[2:])
            run.italic = True
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", line))
        elif line.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s\-|:]+\|\s*$", lines[i + 1].rstrip("\n")
        ):
            header = [c for c in line.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].rstrip("\n").startswith("|"):
                row = [c for c in lines[i].rstrip("\n").strip("|").split("|")]
                rows.append(row)
                i += 1
            add_table_from_md(doc, header, rows)
            continue
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

        i += 1

    doc.save(docx_path)
    print("Wrote", docx_path)


if __name__ == "__main__":
    md_to_docx(
        os.path.join(HERE, "BEGINNERS_GUIDE.md"),
        os.path.join(HERE, "BEGINNERS_GUIDE.docx"),
        "Grasshopper Beginner's Walkthrough — Orthotic Insole",
        font_name="Calibri",
    )
    md_to_docx(
        os.path.join(HERE, "BEGINNERS_GUIDE.hyw.md"),
        os.path.join(HERE, "BEGINNERS_GUIDE_Western_Armenian.docx"),
        "Grasshopper-ի Սկսնակներու Ուղեցոյց — Ուղղագործական Ներբան",
        font_name="Sylfaen",
    )
