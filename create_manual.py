"""Generate the Orthotic Toolkit User Manual as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style tweaks ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0, 70, 140)

# ── Helper functions ──
def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    p.clear()
    run = p.add_run(text)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def add_tip(text):
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    run2 = p.add_run(text)
    run2.italic = True
    return p

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run('WARNING: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    run2 = p.add_run(text)
    run2.italic = True
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 70, 140)
    run2 = p.add_run(text)
    run2.italic = True
    return p

# ══════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ORTHOTIC TOOLKIT v1.0')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0, 70, 140)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Manual & Testing Guide')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('A Professional Insole Design Plugin for Rhinoceros 3D v8')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Plugin Version: 1.0\nRhino Compatibility: 8.0+\nPlatform: Windows')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════

add_heading('Table of Contents', 1)

toc_items = [
    '1. Introduction',
    '2. System Requirements',
    '3. Installation',
    '4. Getting Started — Accessing the Plugin',
    '5. The Complete Workflow — Step by Step',
    '   5.1  Step 1: Select Shoe Last',
    '   5.2  Step 2: Import Foot Scan',
    '   5.3  Step 3: Orient & Extract Plantar Surface',
    '   5.4  Step 4: Generate Insole Outline',
    '   5.5  Step 5: Add Arch Support',
    '   5.6  Step 6: Add Heel Cup',
    '   5.7  Step 7: Add Metatarsal Domes',
    '   5.8  Step 8: Add Posting Wedges',
    '   5.9  Step 9: Apply Thickness Layers',
    '   5.10 Step 10: Validate the Insole',
    '   5.11 Step 11: Export the Insole',
    '6. Panel Reference — All 8 Tabs',
    '7. Command Reference',
    '8. Parameter Quick-Reference Table',
    '9. Layer System',
    '10. Testing Checklist',
    '11. Troubleshooting',
    '12. Tips & Best Practices',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════

add_heading('1. Introduction', 1)
add_para(
    'The Orthotic Toolkit is a professional-grade plugin for Rhinoceros 3D version 8 '
    'that provides a complete workflow for designing custom orthotic insoles. Starting '
    'from a shoe last and a 3D foot scan, the plugin guides you through every step — '
    'from scan import and orientation, through outline generation, arch support, heel '
    'cups, metatarsal domes, posting wedges, thickness layering, validation, and final '
    'export to manufacturing-ready file formats.'
)
add_para(
    'The plugin features a dockable 8-tab panel with intuitive sliders and controls, '
    '14 integrated commands, and support for STL, STEP, OBJ, and 3DM export formats.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════

add_heading('2. System Requirements', 1)

add_bullet('Rhinoceros 3D version 8.0 or later (Windows)')
add_bullet('8 GB RAM recommended (16 GB preferred for large scans)')
add_bullet('Dedicated GPU recommended for smooth viewport performance')
add_bullet('No additional software, compilers, or SDKs required')
add_bullet('Foot scan files in STL, OBJ, or PLY format')
add_bullet('A shoe last model (polysurface) in the Rhino document')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  3. INSTALLATION
# ══════════════════════════════════════════════════════════════════════

add_heading('3. Installation', 1)

add_heading('3.1 Installing the Plugin', 2)
add_numbered('Locate the file OrthoticToolkit_v1.0.rhi in your project folder.')
add_numbered('Double-click the .rhi file. The Rhino Package Installer will open automatically.')
add_numbered('Click "Install" when prompted.')
add_numbered('Restart Rhinoceros 3D to complete the installation.')

add_heading('3.2 Verifying the Installation', 2)
add_numbered('Open Rhino 8.')
add_numbered('Go to Tools > Options > Plug-ins.')
add_numbered('Search for "Orthotic Toolkit" in the plugin list.')
add_numbered('Ensure the checkbox next to it is enabled (checked).')
add_numbered('Type OT_ResetAll in the command line and press Enter. You should see "All state cleared." in the command history.')

add_heading('3.3 Uninstalling', 2)
add_para('To remove the plugin, go to Tools > Options > Plug-ins, find "Orthotic Toolkit", and click "Uninstall". Restart Rhino.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  4. GETTING STARTED
# ══════════════════════════════════════════════════════════════════════

add_heading('4. Getting Started — Accessing the Plugin', 1)

add_para('There are three ways to access the Orthotic Toolkit:', bold=True)

add_heading('4.1 Dockable Panel (Recommended)', 2)
add_numbered('Go to View > Panels.')
add_numbered('Check "Orthotic Toolkit" in the list.')
add_numbered('The panel will appear as a floating window. You can dock it to any side of the Rhino viewport.')
add_para('The panel contains 8 tabs that walk you through the entire workflow, plus a status bar at the bottom.')

add_heading('4.2 Toolbar', 2)
add_numbered('Go to View > Toolbars.')
add_numbered('Check "Orthotic Toolkit".')
add_numbered('A toolbar with 14 command buttons organized in 5 groups (Setup, Scan, Shape, Finish, Utility) will appear.')

add_heading('4.3 Command Line', 2)
add_para(
    'You can type any command directly into the Rhino command line. All commands start '
    'with the prefix OT_ (e.g., OT_SetLast, OT_ImportScan, OT_AddArch).'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  5. COMPLETE WORKFLOW
# ══════════════════════════════════════════════════════════════════════

add_heading('5. The Complete Workflow — Step by Step', 1)
add_para(
    'This section walks you through every step of designing an orthotic insole, from start to finish. '
    'Follow these steps in order for the best results.'
)

# ── Step 1 ──
add_heading('5.1 Step 1: Select Shoe Last (OT_SetLast)', 2)
add_para('Before you can design an insole, you need to tell the plugin which shoe last to use.', bold=True)

add_heading('What You Need:', 3)
add_bullet('A shoe last model (polysurface/Brep) already loaded in your Rhino document.')

add_heading('How to Do It:', 3)
add_numbered('Click the "Select Last" button at the bottom of the Orthotic Toolkit panel, OR type OT_SetLast in the command line.')
add_numbered('The command line will prompt: "Select shoe last polysurface".')
add_numbered('Click on the shoe last in your viewport.')

add_heading('What Happens:', 3)
add_bullet('The plugin automatically detects the sole face (the face pointing downward).')
add_bullet('It projects the sole boundary onto the XY plane to create a green footprint curve.')
add_bullet('It creates an inverse sole surface (offset inward by the cover thickness).')
add_bullet('The status bar updates to show the selected last name.')

add_heading('What You Should See:', 3)
add_bullet('A green footprint curve on the OT_Preview layer.')
add_bullet('The status bar at the bottom of the panel shows the last name.')

add_tip('Make sure your shoe last is a closed polysurface (solid). If it is an open surface, the sole detection may not work correctly.')

# ── Step 2 ──
add_heading('5.2 Step 2: Import Foot Scan (OT_ImportScan)', 2)
add_para('Import a 3D scan of the patient\'s foot.', bold=True)

add_heading('Supported File Formats:', 3)
add_bullet('STL (binary or ASCII)')
add_bullet('OBJ (Wavefront)')
add_bullet('PLY (Polygon File Format)')

add_heading('How to Do It:', 3)
add_numbered('Go to the "Foot Scan" tab (Tab 1) in the panel.')
add_numbered('Click "Import Scan...".')
add_numbered('A file dialog will open. Navigate to your scan file and click Open.')
add_numbered('After the file imports, you will be prompted to select the mesh in the viewport. Click on the imported mesh.')

add_heading('What You Should See:', 3)
add_bullet('The foot scan mesh appears on the OT_FootScan layer (light blue).')
add_bullet('The status bar shows the imported filename.')

add_warning('Typical foot scan meshes have 40,000–100,000+ vertices. Very large scans may take a moment to import.')

# ── Step 3 ──
add_heading('5.3 Step 3: Orient & Extract Plantar Surface', 2)

add_heading('3a. Auto-Orient the Scan (OT_OrientScan)', 3)
add_para('The scan needs to be oriented so the sole of the foot faces downward (-Z direction).')
add_numbered('In the "Foot Scan" tab, optionally select an Orientation Preset from the dropdown (Scanner Default, Rotated 90 X, Rotated 180 Z).')
add_numbered('Click "Auto-Orient".')
add_numbered('The plugin analyzes the mesh normals and rotates the scan so the plantar surface faces -Z.')
add_numbered('The command line reports the rotation angle applied.')

add_note('If the scan is already correctly oriented (within 10 degrees), it will remain unchanged.')

add_heading('3b. Extract Plantar Surface (OT_ExtractPlantar)', 3)
add_para('This creates a smooth NURBS surface from the bottom of the foot scan.')
add_numbered('Adjust the "Smoothing Passes" slider (0–5). Default is 2. Higher values = smoother surface.')
add_numbered('Click "Extract Plantar Surface".')
add_numbered('The plugin shoots a grid of rays (40x20) upward through the scan to sample the plantar surface, then fits a smooth NURBS surface through the points.')

add_heading('What You Should See:', 3)
add_bullet('A yellow NURBS surface appears on the OT_PlantarSurface layer.')
add_bullet('The status label shows point count and surface degree.')

add_tip('If the extraction reports fewer than 200 intersection points, try adjusting the scan orientation or reducing smoothing passes.')

# ── Step 4 ──
add_heading('5.4 Step 4: Generate Insole Outline (OT_GenerateOutline)', 2)
add_para('Create the 2D perimeter of the insole and the initial flat insole shape.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Outline" tab (Tab 2).')
add_numbered('Adjust the parameters:')
add_bullet('Perimeter Offset (0–10mm, default 2mm): How far inward from the footprint edge the insole boundary sits.', level=1)
add_bullet('Toe Extension (-5 to +10mm, default 0mm): Extend or shorten the toe area.', level=1)
add_bullet('Heel Extension (-5 to +10mm, default 0mm): Extend or shorten the heel area.', level=1)
add_numbered('Click "Generate Outline".')

add_heading('What You Should See:', 3)
add_bullet('An orange outline curve on the OT_Outline layer.')
add_bullet('A blue flat insole solid (Brep) on the OT_Insole layer.')

add_note('The insole Brep is extruded downward with a total height equal to cover + shell + base thickness (default: 2 + 3 + 5 = 10mm).')

# ── Step 5 ──
add_heading('5.5 Step 5: Add Arch Support (OT_AddArch)', 2)
add_para('Add a medial arch support to the insole.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Arch" tab (Tab 3).')
add_numbered('Adjust the parameters:')
add_bullet('Arch Height (0–20mm, default 10mm): How tall the arch peak is.', level=1)
add_bullet('Apex Position (0–100%, default 50%): Where along the length of the insole the arch peaks. 0% = heel end, 100% = toe end.', level=1)
add_bullet('Width (10–40mm, default 20mm): How wide the arch support is.', level=1)
add_bullet('Blend Radius (0–5mm, default 3mm): How smoothly the arch blends into the insole.', level=1)
add_numbered('Click "Apply Arch".')

add_heading('What You Should See:', 3)
add_bullet('The insole on the OT_Insole layer updates with a raised arch on the medial (inner) side.')

add_tip('Start with default values and adjust incrementally. Very tall arches (>15mm) may cause boolean union failures.')

# ── Step 6 ──
add_heading('5.6 Step 6: Add Heel Cup (OT_AddHeelCup)', 2)
add_para('Add a U-shaped heel cup to cradle the heel.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Heel Cup" tab (Tab 4).')
add_numbered('Adjust the parameters:')
add_bullet('Cup Depth (0–20mm, default 12mm): How deep the heel cup walls are.', level=1)
add_bullet('Posterior Angle (70–110°, default 90°): The tilt angle of the back wall.', level=1)
add_bullet('Lateral Flare (0–30°, default 10°): Outward flare of the outer wall.', level=1)
add_bullet('Medial Flare (0–30°, default 10°): Outward flare of the inner wall.', level=1)
add_bullet('Cup Width % (50–100%, default 100%): How much of the heel width is covered.', level=1)
add_numbered('Click "Apply Heel Cup".')

add_heading('What You Should See:', 3)
add_bullet('The insole updates with raised walls around the heel area.')

# ── Step 7 ──
add_heading('5.7 Step 7: Add Metatarsal Domes (OT_AddMetDome)', 2)
add_para('Add dome-shaped pads in the forefoot area for pressure redistribution.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Forefoot" tab (Tab 5).')
add_numbered('Set the Dome Count (1–5, default 1).')
add_numbered('Click "Apply Met Domes".')

add_heading('What You Should See:', 3)
add_bullet('Dome-shaped bumps appear on the forefoot area of the insole.')

add_note('Dome positions are automatically calculated based on the insole geometry. Default dome height is 5mm and diameter is 10mm.')

# ── Step 8 ──
add_heading('5.8 Step 8: Add Posting Wedges (OT_AddPosting)', 2)
add_para('Add angular wedges for pronation/supination control.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Posting" tab (Tab 6).')
add_numbered('Adjust the parameters:')
add_bullet('Rearfoot Medial (-15 to +15°, default 0°): Tilts the rearfoot toward the medial side.', level=1)
add_bullet('Rearfoot Lateral (-15 to +15°, default 0°): Tilts the rearfoot toward the lateral side.', level=1)
add_bullet('Forefoot Medial (-15 to +15°, default 0°): Tilts the forefoot toward the medial side.', level=1)
add_bullet('Forefoot Lateral (-15 to +15°, default 0°): Tilts the forefoot toward the lateral side.', level=1)
add_bullet('Split % (10–90%, default 50%): Where the rearfoot/forefoot boundary is.', level=1)
add_numbered('Click "Apply Posting".')

add_heading('What You Should See:', 3)
add_bullet('The insole surfaces tilt according to the angles you specified.')

add_tip('Wedges with angles less than 0.01° are automatically skipped (no change).')

# ── Step 9 ──
add_heading('5.9 Step 9: Apply Thickness Layers (OT_SetThickness)', 2)
add_para('Split the insole into three structural layers and set their thicknesses.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Thickness" tab (Tab 7).')
add_numbered('Adjust the layer thicknesses:')
add_bullet('Cover (0–5mm, default 2mm): The top surface layer.', level=1)
add_bullet('Shell (0–10mm, default 3mm): The structural middle layer.', level=1)
add_bullet('Base (0–10mm, default 5mm): The bottom layer.', level=1)
add_numbered('The "Total Thickness" label updates automatically as you adjust.')
add_numbered('Click "Apply Thickness".')

add_heading('What You Should See:', 3)
add_bullet('The insole is split into three distinct layers.')
add_bullet('If any area is thinner than 2mm, a red point cloud appears on the OT_Warnings layer and a warning dialog is shown identifying the thin regions (heel, arch, or forefoot).')

add_warning('Pay attention to minimum thickness warnings. Areas thinner than 2mm may be structurally weak.')

# ── Step 10 ──
add_heading('5.10 Step 10: Validate the Insole (OT_ValidateInsole)', 2)
add_para('Run a comprehensive validation check before exporting.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('In the "Thickness" tab (Tab 7), click "Run Validation".')
add_numbered('A dialog will appear showing the validation results.')

add_heading('Validation Checks:', 3)
add_bullet('IsValid: Checks the Brep topology. Attempts auto-repair if invalid.')
add_bullet('IsSolid: Confirms the Brep is a closed solid with no gaps or holes.')
add_bullet('Minimum Thickness: Samples 200 random points and checks that all are >= 2mm thick.')
add_bullet('Overhangs: Counts faces that overhang more than 45° (relevant for 3D printing).')
add_bullet('Overall Status: Reports either "READY FOR EXPORT" or "ISSUES FOUND".')

add_heading('What You Want to See:', 3)
add_para('All checks should pass, and the overall status should read "READY FOR EXPORT".', bold=True)

# ── Step 11 ──
add_heading('5.11 Step 11: Export the Insole (OT_ExportInsole)', 2)
add_para('Export the final insole to a manufacturing-ready file.', bold=True)

add_heading('How to Do It:', 3)
add_numbered('Go to the "Export" tab (Tab 8).')
add_numbered('Choose your export settings:')
add_bullet('Format: STL (3D printing/CAM), STEP (CAD/engineering), OBJ (3D modeling), or 3DM (Rhino native).', level=1)
add_bullet('Mesh Resolution (for STL/OBJ): Draft (0.5mm), Standard (0.2mm), Fine (0.1mm, recommended), or Ultra (0.05mm).', level=1)
add_bullet('Export by Layer: If checked, creates separate files for cover, shell, and base layers (e.g., insole_cover.stl, insole_shell.stl, insole_base.stl).', level=1)
add_bullet('Include Rocker: If checked, includes a rocker contact outline curve in the export.', level=1)
add_numbered('Click "Export...".')
add_numbered('Choose a save location and filename in the file dialog.')

add_heading('What You Should See:', 3)
add_bullet('A confirmation message in the command history.')
add_bullet('The exported file(s) at your chosen location.')

add_tip('For 3D printing, use STL format with Fine (0.1mm) resolution. For CAD interchange, use STEP format.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  6. PANEL REFERENCE
# ══════════════════════════════════════════════════════════════════════

add_heading('6. Panel Reference — All 8 Tabs', 1)

tabs = [
    ('Tab 1: Foot Scan', [
        '"Import Scan..." button — opens file dialog for STL/OBJ/PLY',
        'Orientation Preset dropdown — Scanner Default, Rotated 90 X, Rotated 180 Z',
        '"Auto-Orient" button — automatically orients scan so sole faces down',
        'Smoothing Passes slider (0–5, default 2)',
        '"Extract Plantar Surface" button — creates NURBS surface from scan',
        'Status label — shows extraction results',
    ]),
    ('Tab 2: Outline', [
        'Perimeter Offset slider + spinbox (0–10mm)',
        'Toe Extension slider + spinbox (-5 to +10mm)',
        'Heel Extension slider + spinbox (-5 to +10mm)',
        '"Generate Outline" button',
    ]),
    ('Tab 3: Arch', [
        'Arch Height slider + spinbox (0–20mm)',
        'Apex Position slider + spinbox (0–100%)',
        'Width slider + spinbox (10–40mm)',
        'Blend Radius slider + spinbox (0–5mm)',
        '"Apply Arch" button',
    ]),
    ('Tab 4: Heel Cup', [
        'Cup Depth slider + spinbox (0–20mm)',
        'Posterior Angle slider + spinbox (70–110°)',
        'Lateral Flare slider + spinbox (0–30°)',
        'Medial Flare slider + spinbox (0–30°)',
        'Cup Width % slider + spinbox (50–100%)',
        '"Apply Heel Cup" button',
    ]),
    ('Tab 5: Forefoot', [
        'Dome Count spinbox (1–5)',
        '"Apply Met Domes" button',
    ]),
    ('Tab 6: Posting', [
        'Rearfoot Medial slider + spinbox (-15 to +15°)',
        'Rearfoot Lateral slider + spinbox (-15 to +15°)',
        'Forefoot Medial slider + spinbox (-15 to +15°)',
        'Forefoot Lateral slider + spinbox (-15 to +15°)',
        'Split % slider + spinbox (10–90%)',
        '"Apply Posting" button',
    ]),
    ('Tab 7: Thickness', [
        'Cover Thickness slider + spinbox (0–5mm)',
        'Shell Thickness slider + spinbox (0–10mm)',
        'Base Thickness slider + spinbox (0–10mm)',
        'Total Thickness label (auto-calculated, read-only)',
        '"Apply Thickness" button',
        '"Run Validation" button',
    ]),
    ('Tab 8: Export', [
        'Format dropdown — STL, STEP, OBJ, 3DM',
        'Mesh Resolution dropdown — Draft, Standard, Fine, Ultra',
        'Export by Layer checkbox',
        'Include Rocker checkbox',
        '"Export..." button',
    ]),
]

for tab_name, controls in tabs:
    add_heading(tab_name, 2)
    for ctrl in controls:
        add_bullet(ctrl)

add_heading('Status Bar (Bottom of Panel)', 2)
add_bullet('"Select Last" button — runs OT_SetLast')
add_bullet('Last name label — shows currently selected shoe last')
add_bullet('Scan filename label — shows currently imported scan file')
add_bullet('"Reset All" button — clears all state and starts fresh')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  7. COMMAND REFERENCE
# ══════════════════════════════════════════════════════════════════════

add_heading('7. Command Reference', 1)
add_para('All commands can be typed directly into the Rhino command line. They all use the OT_ prefix.')

add_table(
    ['Command', 'Category', 'Description'],
    [
        ['OT_SetLast', 'Setup', 'Select shoe last and generate footprint'],
        ['OT_ImportScan', 'Scan', 'Import foot scan mesh (STL/OBJ/PLY)'],
        ['OT_OrientScan', 'Scan', 'Auto-orient scan so sole faces down'],
        ['OT_ExtractPlantar', 'Scan', 'Extract NURBS plantar surface from scan'],
        ['OT_GenerateOutline', 'Base', 'Create insole outline and flat insole shape'],
        ['OT_AddArch', 'Feature', 'Add medial arch support'],
        ['OT_AddHeelCup', 'Feature', 'Add U-shaped heel cup'],
        ['OT_AddMetDome', 'Feature', 'Add metatarsal dome pads'],
        ['OT_AddPosting', 'Feature', 'Add posting wedges for pronation/supination'],
        ['OT_SetThickness', 'Layer', 'Split insole into cover/shell/base layers'],
        ['OT_ValidateInsole', 'Validate', 'Run full validation checks'],
        ['OT_RockerOutline', 'Utility', 'Generate rocker contact outline curve'],
        ['OT_ExportInsole', 'Export', 'Validate and export insole to file'],
        ['OT_ResetAll', 'Utility', 'Clear all state and start fresh'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  8. PARAMETER QUICK-REFERENCE
# ══════════════════════════════════════════════════════════════════════

add_heading('8. Parameter Quick-Reference Table', 1)

add_table(
    ['Feature', 'Parameter', 'Default', 'Range', 'Unit'],
    [
        ['Outline', 'Perimeter Offset', '2.0', '0–10', 'mm'],
        ['Outline', 'Toe Extension', '0.0', '-5 to +10', 'mm'],
        ['Outline', 'Heel Extension', '0.0', '-5 to +10', 'mm'],
        ['Arch', 'Height', '10.0', '0–20', 'mm'],
        ['Arch', 'Apex Position', '50', '0–100', '%'],
        ['Arch', 'Width', '20.0', '10–40', 'mm'],
        ['Arch', 'Blend Radius', '3.0', '0–5', 'mm'],
        ['Heel Cup', 'Depth', '12.0', '0–20', 'mm'],
        ['Heel Cup', 'Posterior Angle', '90.0', '70–110', '°'],
        ['Heel Cup', 'Lateral Flare', '10.0', '0–30', '°'],
        ['Heel Cup', 'Medial Flare', '10.0', '0–30', '°'],
        ['Heel Cup', 'Width %', '100', '50–100', '%'],
        ['Met Dome', 'Count', '1', '1–5', 'count'],
        ['Posting', 'RF Medial', '0.0', '-15 to +15', '°'],
        ['Posting', 'RF Lateral', '0.0', '-15 to +15', '°'],
        ['Posting', 'FF Medial', '0.0', '-15 to +15', '°'],
        ['Posting', 'FF Lateral', '0.0', '-15 to +15', '°'],
        ['Posting', 'Split %', '50', '10–90', '%'],
        ['Thickness', 'Cover', '2.0', '0–5', 'mm'],
        ['Thickness', 'Shell', '3.0', '0–10', 'mm'],
        ['Thickness', 'Base', '5.0', '0–10', 'mm'],
        ['Scan', 'Smoothing Passes', '2', '0–5', 'count'],
        ['Export', 'Mesh Tolerance', '0.1 (Fine)', '0.05–0.5', 'mm'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  9. LAYER SYSTEM
# ══════════════════════════════════════════════════════════════════════

add_heading('9. Layer System', 1)
add_para(
    'The plugin automatically creates and manages layers for organizing geometry. '
    'Each step places its output on a specific layer with a designated color.'
)

add_table(
    ['Layer Name', 'Color', 'Contents'],
    [
        ['OT_Preview', 'Green', 'Footprint curve from shoe last'],
        ['OT_FootScan', 'Light Blue', 'Imported foot scan mesh'],
        ['OT_PlantarSurface', 'Yellow', 'Extracted NURBS plantar surface'],
        ['OT_Outline', 'Orange', 'Insole perimeter outline curve'],
        ['OT_Insole', 'Blue', 'Main insole solid (updated by all shape tools)'],
        ['OT_Warnings', 'Red', 'Thin-area warning point clouds'],
        ['OT_RockerContact', 'Purple', 'Rocker contact outline curve'],
    ]
)

add_tip('You can toggle layer visibility in Rhino\'s Layers panel to show/hide different elements of your design.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  10. TESTING CHECKLIST
# ══════════════════════════════════════════════════════════════════════

add_heading('10. Testing Checklist', 1)
add_para(
    'Use this checklist to systematically test every feature of the plugin. '
    'Work through each item and verify the expected result.'
)

tests = [
    ('Installation & Startup', [
        ('Double-click OrthoticToolkit_v1.0.rhi and install', 'Rhino Package Installer opens; install completes'),
        ('Restart Rhino 8', 'No errors on startup'),
        ('View > Panels > Orthotic Toolkit', 'Dockable panel appears with 8 tabs'),
        ('Type OT_ResetAll', '"All state cleared." in command history'),
    ]),
    ('Step 1: Select Last', [
        ('Load a shoe last model into your Rhino file', 'Polysurface visible in viewport'),
        ('Click "Select Last" or type OT_SetLast', 'Prompted to select polysurface'),
        ('Click on the shoe last', 'Green footprint curve appears; status bar shows last name'),
    ]),
    ('Step 2: Import Scan', [
        ('Click "Import Scan..." in Tab 1', 'File dialog opens with STL/OBJ/PLY filter'),
        ('Select a foot scan file', 'Mesh imports; prompted to select in viewport'),
        ('Click on the imported mesh', 'Mesh appears light blue on OT_FootScan layer; status bar shows filename'),
    ]),
    ('Step 3: Orient & Extract', [
        ('Click "Auto-Orient" in Tab 1', 'Mesh rotates; rotation angle reported in command line'),
        ('Adjust smoothing slider to 2', 'Slider label updates to "2"'),
        ('Click "Extract Plantar Surface"', 'Yellow NURBS surface on OT_PlantarSurface layer; status shows point count'),
    ]),
    ('Step 4: Generate Outline', [
        ('Set Perimeter Offset to 2mm in Tab 2', 'Value shows 2.0'),
        ('Click "Generate Outline"', 'Orange outline curve + blue insole solid appear'),
    ]),
    ('Step 5: Add Arch', [
        ('Set Height=10, Apex=50%, Width=20, Radius=3 in Tab 3', 'Values update in sliders'),
        ('Click "Apply Arch"', 'Insole updates with visible arch on medial side'),
    ]),
    ('Step 6: Add Heel Cup', [
        ('Set Depth=12, Angle=90, Flares=10/10, Width=100% in Tab 4', 'Values update'),
        ('Click "Apply Heel Cup"', 'Insole updates with heel cup walls'),
    ]),
    ('Step 7: Add Met Domes', [
        ('Set Dome Count=1 in Tab 5', 'Value shows 1'),
        ('Click "Apply Met Domes"', 'Dome bump visible on forefoot area'),
    ]),
    ('Step 8: Add Posting', [
        ('Set RF Medial=4, others=0, Split=50% in Tab 6', 'Values update'),
        ('Click "Apply Posting"', 'Rearfoot area tilts medially'),
    ]),
    ('Step 9: Thickness', [
        ('Set Cover=2, Shell=3, Base=5 in Tab 7', 'Total shows 10mm'),
        ('Click "Apply Thickness"', 'Insole splits into layers; check for thin-area warnings'),
    ]),
    ('Step 10: Validate', [
        ('Click "Run Validation" in Tab 7', 'Dialog shows all 5 checks'),
        ('Review results', 'Ideally shows "READY FOR EXPORT"'),
    ]),
    ('Step 11: Export', [
        ('Select STL format, Fine resolution in Tab 8', 'Dropdowns update'),
        ('Click "Export..."', 'File save dialog opens'),
        ('Save file', 'STL file created at chosen location'),
        ('Try STEP format export', 'STP file created successfully'),
        ('Check "Export by Layer", export again', 'Three separate files created (_cover, _shell, _base)'),
    ]),
    ('Utility', [
        ('Click "Reset All"', 'All state cleared; panel labels reset; preview objects removed'),
        ('Repeat workflow from Step 1', 'Full workflow works again from scratch'),
    ]),
]

for section_name, items in tests:
    add_heading(section_name, 2)
    tbl = doc.add_table(rows=1 + len(items), cols=3)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Action', 'Expected Result', 'Pass?']):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for r_idx, (action, expected) in enumerate(items):
        tbl.rows[r_idx + 1].cells[0].text = action
        tbl.rows[r_idx + 1].cells[1].text = expected
        tbl.rows[r_idx + 1].cells[2].text = '[ ]'
        for c in range(3):
            for p in tbl.rows[r_idx + 1].cells[c].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  11. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════

add_heading('11. Troubleshooting', 1)

issues = [
    ('Panel does not appear in View > Panels',
     'Go to Tools > Options > Plug-ins. Find "Orthotic Toolkit" and make sure it is enabled (checked). Restart Rhino.'),
    ('"Unknown command" when typing OT_ commands',
     'The plugin may not have loaded. Check the Plug-ins dialog and restart Rhino.'),
    ('Boolean union fails (arch, heel cup, or posting)',
     'Try less extreme parameter values. Very tall arches (>15mm) or very deep heel cups can cause geometry failures. Reduce the values and try again.'),
    ('Export produces an empty file',
     'Run OT_ValidateInsole first. Ensure the status is "READY FOR EXPORT". If not, fix the reported issues.'),
    ('Scan import shows no mesh',
     'Verify the file is a valid STL/OBJ/PLY. After import, you may need to click on the mesh in the viewport to select it.'),
    ('Plantar extraction gets too few points',
     'The scan may not be correctly oriented. Run OT_OrientScan first. Also try reducing smoothing passes.'),
    ('Red warning points appear after thickness',
     'Some areas of the insole are thinner than 2mm. Review the thin regions shown in the warning dialog and adjust your design parameters.'),
    ('Plugin crashes or shows an error dialog',
     'The global exception handler will show details. Note the error message, restart Rhino, and try the operation again with different parameters.'),
]

for problem, solution in issues:
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    run2 = p.add_run(problem)

    p2 = doc.add_paragraph()
    run3 = p2.add_run('Solution: ')
    run3.bold = True
    run3.font.color.rgb = RGBColor(0, 128, 0)
    run4 = p2.add_run(solution)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
#  12. TIPS & BEST PRACTICES
# ══════════════════════════════════════════════════════════════════════

add_heading('12. Tips & Best Practices', 1)

tips = [
    'Always start with OT_SetLast before any other command. Most commands require a shoe last to be selected.',
    'Save your Rhino file frequently. The plugin state does not persist after closing Rhino.',
    'Use the default parameter values first, then adjust incrementally to see their effect.',
    'If a boolean operation fails, undo (Ctrl+Z) and try with smaller parameter values.',
    'For 3D printing, always run validation and export as STL with Fine (0.1mm) resolution.',
    'For CAD interchange with engineering software, use STEP format.',
    'Toggle layer visibility (Rhino Layers panel) to inspect individual design elements.',
    'The "Export by Layer" option is useful when each layer will be manufactured from a different material.',
    'Use OT_ResetAll to start a completely fresh design without restarting Rhino.',
    'Keep foot scan meshes under 100K vertices for best performance. Decimate larger scans in your scanning software before import.',
    'The smoothing slider (0–5 passes) on the Foot Scan tab lets you control plantar surface smoothness. 2 passes is a good default.',
    'Use the rocker outline (Include Rocker checkbox) if your manufacturing process requires a rocker-bottom contact reference.',
]

for t in tips:
    add_bullet(t)

# ── Save ──
output_path = '/home/user/Rhino-Orthotic-Toolkit/Orthotic_Toolkit_User_Manual.docx'
doc.save(output_path)
print(f'Manual saved to: {output_path}')
